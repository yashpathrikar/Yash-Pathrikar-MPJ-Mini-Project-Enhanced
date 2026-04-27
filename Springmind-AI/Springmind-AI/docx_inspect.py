from docx import Document

path = r"C:\Users\ASUS\Documents\Springmind-AI\Ticket_Management_Enhancement_Report_SLA_Only.docx"
doc = Document(path)

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        print(f"P{i}: {text}")

for ti, table in enumerate(doc.tables):
    print(f"TABLE {ti}")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.replace("\n", " | ").strip() for cell in row.cells]
        print(f"  R{ri}: {cells}")
