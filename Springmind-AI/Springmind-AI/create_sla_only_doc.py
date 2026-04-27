from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\Users\ASUS\Documents\Springmind-AI\Ticket_Management_Enhancement_Report_SLA_Only.docx"

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(7)
styles["Normal"].paragraph_format.line_spacing = 1.08

for style_name, size, color in [
    ("Title", 22, RGBColor(31, 78, 121)),
    ("Heading 1", 15, RGBColor(31, 78, 121)),
    ("Heading 2", 12, RGBColor(68, 68, 68)),
]:
    style = styles[style_name]
    style.font.name = "Aptos Display" if style_name == "Title" else "Aptos"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color


def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run("- " + item)
        run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(3)


# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MINIPROJECT ENHANCEMENT REPORT")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(31, 78, 121)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Ticket Management System - Enhancement Report")
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SLA-Based Ticket Prioritization Logic")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()
meta = [
    ("Student Name:", "Pratiyush Sharma"),
    ("Enrollment / ID:", "1032232854"),
    ("Course / Department:", "Computer Science Engineering / Computer Engineering & Technology"),
    ("Institution Name:", "Dr. Vishwanath Karad MIT World Peace University"),
    ("Guide Name:", "Dr. Suhas Joshi"),
    ("Submission Date:", "27 April, 2026"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(value)

doc.add_page_break()

# Abstract - preserved exactly as provided in the original document.
add_heading("ABSTRACT", 1)
abstract_paragraphs = [
    "This report documents the proposed backend enhancements for the Ticket Management System. The enhancement scope focuses on strengthening account security, improving operational ticket prioritization through SLA-aware logic, introducing automated progress reporting through scheduled emails with generated PDF attachments, converting numerical ticket data into graphical visualizations, and notifying the Admin dashboard when new tickets are raised or ticket statuses change.",
    "The enhancements span five major domains: (1) Email-based Two-Factor Authentication (2FA), CAPTCHA verification, and secure forgot-password recovery for both Admin and User accounts; (2) SLA-based ticket prioritization that ensures tickets nearing breach are resolved before newer non-critical tickets; (3) Automated progress emails every 10 minutes with a PDF report containing ticket IDs, statuses, progress updates, and timestamps; (4) Graphical data visualization that converts numerical ticket data into charts for better understanding; and (5) Admin dashboard notifications for newly raised tickets and ticket status updates.",
    "The implementation is designed to extend the current backend without breaking existing APIs, database behavior, or the established UI/UX. Required database changes are additive, email, PDF, visualization, and notification services are modularized, and background scheduling is introduced with safeguards to avoid duplicate or skipped notifications.",
]
for text in abstract_paragraphs:
    doc.add_paragraph(text)

add_heading("ENHANCEMENT OVERVIEW", 1)
add_heading("Background", 2)
doc.add_paragraph(
    "The ticket system already supports ticket creation, ticket status management, and user/admin workflows. "
    "The SLA-based prioritization enhancement focuses only on improving operational ticket handling so that work queues are ordered by deadline risk instead of simple newest-first behavior."
)

add_heading("Enhancement Goal", 2)
add_bullets([
    "Prioritize tickets by SLA urgency so tickets close to breach are handled first.",
    "Prevent newer non-critical tickets from overriding older or more urgent SLA-critical tickets.",
    "Keep the existing ticket creation flow unchanged and apply prioritization only during listing, assignment, or resolution queue handling.",
    "Maintain backward compatibility by adding optional SLA fields only where useful.",
])

add_heading("Enhancement Summary Table", 2)
doc.add_paragraph(
    "SLA-Based Prioritization: Deadline-aware sorting, SLA-critical escalation, and unchanged ticket creation flow."
)

add_heading("ENHANCEMENT: SLA-BASED TICKET PRIORITIZATION LOGIC", 1)
doc.add_paragraph(
    "The ticket handling logic is enhanced so that SLA-critical tickets are resolved before newer tickets. "
    "This ensures operational fairness is based on deadline risk rather than simple creation order alone."
)

add_heading("Prioritization Rule", 2)
add_bullets([
    "Each ticket is evaluated using its SLA deadline or expected closure time.",
    "Tickets close to SLA breach are assigned the highest resolution priority.",
    "Newer tickets must not override older tickets that are near SLA breach.",
    "Existing ticket creation flow remains unchanged; prioritization is applied when fetching or assigning work.",
])

add_heading("Sorting Logic", 2)
add_bullets([
    f"{condition}: {behavior}"
    for condition, behavior in [
    ("Near SLA breach", "Resolved first regardless of ticket creation time."),
    ("Already breached", "Escalated to the top and flagged for immediate handling."),
    ("Normal active ticket", "Sorted by configured priority and creation time after SLA-critical tickets."),
    ("Closed/resolved ticket", "Excluded from active resolution queue unless reporting requires it."),
    ]
])

add_heading("Compatibility Requirements", 2)
add_bullets([
    "No change is made to the ticket creation request payload unless additional SLA metadata is required.",
    "Database changes are additive only, such as SLA deadline, SLA status, or last prioritization timestamp fields.",
    "Existing APIs continue to return compatible responses, with optional SLA fields added where useful.",
    "Business logic is isolated in a ticket prioritization service for clean testing and future modification.",
])

add_heading("TECHNICAL IMPLEMENTATION PLAN", 1)
add_heading("Backend Module", 2)
add_bullets([
    "Ticket Prioritization Service: Apply SLA-based ordering before assignment or resolution queue display.",
    "Ticket Service Integration: Use SLA deadline sorting when fetching active queues without changing ticket creation.",
    "Repository Query Layer: Sort unresolved tickets by nearest SLA deadline, then fallback priority and creation time.",
])

add_heading("Required Database Changes", 2)
add_bullets([
    "tickets.sla_deadline: Represent the expected closure time used for SLA prioritization.",
    "tickets.sla_status: Optional status such as normal, near_breach, breached, or resolved.",
    "tickets.last_prioritized_at: Optional audit timestamp for future queue analysis.",
])

add_heading("Clean Code and Compatibility Notes", 2)
add_bullets([
    "Add clear comments only around SLA ordering logic where the business rule needs explanation.",
    "Keep controller contracts backward compatible and add optional response fields rather than removing existing fields.",
    "Use migrations for all additive schema changes so environments can be updated predictably.",
    "Keep SLA prioritization logic separate from ticket creation so current user flows remain stable.",
])

add_heading("OUTCOME & RESULTS", 1)
add_heading("Expected Feature Delivered", 2)
add_bullets([
    "SLA Prioritization: Near-breach tickets are sorted above newer non-critical tickets.",
    "Resolution Queue Discipline: Agents see SLA-critical work first, reducing avoidable SLA breaches.",
    "Backward Compatibility: Existing APIs, UI flows, and ticket creation behavior continue to work after deployment.",
])

add_heading("Validation Checklist", 2)
add_bullets([
    "Tickets close to SLA breach remain ahead of newer tickets in resolution queues.",
    "Already breached tickets are escalated and clearly flagged.",
    "Closed or resolved tickets do not interfere with active SLA resolution ordering.",
    "Existing ticket creation behavior continues to work without request payload changes.",
])

add_heading("CONCLUSION", 1)
doc.add_paragraph(
    "This enhancement strengthens the Ticket Management System by introducing SLA-aware ticket handling. "
    "The improved queue ordering ensures tickets nearing breach are resolved before newer non-critical tickets, helping agents respond according to business urgency and deadline risk."
)
doc.add_paragraph(
    "The proposed design follows clean code practices by separating ticket prioritization into focused backend logic. "
    "Database updates remain additive and migration-driven, preserving backward compatibility while enabling reliable SLA-based queue behavior."
)

add_heading("Learning Outcomes / Implementation Focus", 2)
add_bullets([
    "SLA-driven queue ordering that aligns system behavior with business urgency.",
    "Backward-compatible backend extension using modular services and additive migrations.",
    "Priority handling that prevents newer tickets from overriding near-breach tickets.",
])

doc.save(OUT)
print(OUT)
